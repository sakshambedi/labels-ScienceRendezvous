from pathlib import Path
import pandas as pd
import numpy as np
import os
from docx import Document        
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pandas.core import frame
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


document = Document()
shipping_styles = document.styles
shipping_charstyle = shipping_styles.add_style('shippingStyle', WD_STYLE_TYPE.CHARACTER)
shipping_font = shipping_charstyle.font
shipping_font.name = 'Calibri (Body)'
shipping_font.size = Pt(28)

# rcpnt_data_array contains the following information in the following order :
# rcpnt_name,rcpnt_title, rcpnt_school_name, rcpnt_adrress , rcpnt_city, rcpnt_postal_code
def edit_docx(rcpnt_data_array):
        
        reference_file = open(Path(os.getcwd(), 'assets', 'university_shipping_title.txt'), "r")
        
        
        paragraph = document.add_paragraph(reference_file.read())
        style =  document.styles['Normal']
        font = style.font
        font.name = 'Calibri (Body)'
        font.size = Pt(18)
        paragraph.style = document.styles['Normal']

        s  = document.add_paragraph('\n')
        shipping_paragraph = s.add_run(('To: Attn: ', rcpnt_data_array[0],', ', rcpnt_data_array[1], '\n\n',rcpnt_data_array[2],'\n',rcpnt_data_array[3],'\n',rcpnt_data_array[4] ,',',rcpnt_data_array[5]),
                                        style = 'shippingStyle').bold = True
        s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        document.add_page_break()
        

def read_xlsx():
    working_sheet_name = 'Sheet1'
    loc_file = Path(os.getcwd(),'assets','Shipping.xlsx')
    
    working_excel_file = pd.read_excel(loc_file, sheet_name = working_sheet_name)
    df = pd.DataFrame(working_excel_file,columns=['School Contact', 'Position Title', 
                                                'School Name', 'Delivery Address', 'City','Postal Code'])
    # print(df)
    for each_row in df.index:
        # rcpnt == recipient
        rcpnt_name = df['School Contact'][each_row]
        rcpnt_title = df['Position Title'][each_row]
        rcpnt_school_name =  df['School Name'][each_row]
        rcpnt_adrress = df['Delivery Address'][each_row]
        rcpnt_city = df['City'][each_row]
        rcpnt_postal_code = df['Postal Code'][each_row]
        arr = np.array([rcpnt_name,rcpnt_title, rcpnt_school_name, rcpnt_adrress , rcpnt_city, rcpnt_postal_code])
        edit_docx(arr)

    document.save(Path(os.getcwd(), 'assets', 'Shipping-Labels.docx'))

def main():
    read_xlsx()

if __name__ == '__main__':
    main()
